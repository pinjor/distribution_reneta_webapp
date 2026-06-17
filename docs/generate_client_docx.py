"""Convert SWIFT_DISTRIBUTION_HUB_CLIENT_DOCUMENTATION.md to a Word document with proper tables."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = Path(__file__).resolve().parent
MD_FILE = ROOT / "SWIFT_DISTRIBUTION_HUB_CLIENT_DOCUMENTATION.md"
DOCX_FILE = ROOT / "SWIFT_DISTRIBUTION_HUB_CLIENT_DOCUMENTATION.docx"


def strip_md_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text.strip()


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|[\s\-:|]+\|$", line.strip()))


def parse_table_row(line: str) -> list[str]:
    cells = [strip_md_inline(c.strip()) for c in line.strip().strip("|").split("|")]
    return cells


def set_cell_shading(cell, fill: str = "D9E2F3") -> None:
    from docx.oxml import OxmlElement

    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return

    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Column width hints by column count
    page_width = Inches(6.5)
    header_text = " ".join(rows[0]).lower() if rows else ""
    if cols == 2:
        widths = [Inches(2.0), Inches(4.5)]
    elif cols == 3:
        if "location" in header_text and "purpose" in header_text:
            widths = [Inches(2.0), Inches(2.3), Inches(2.2)]
        elif "owner" in header_text or "dependency" in header_text:
            widths = [Inches(1.6), Inches(2.4), Inches(2.5)]
        else:
            widths = [Inches(1.8), Inches(2.2), Inches(2.5)]
    elif cols == 4:
        widths = [Inches(1.4), Inches(1.6), Inches(1.8), Inches(1.7)]
    elif cols == 5:
        widths = [Inches(1.2), Inches(1.3), Inches(1.3), Inches(1.4), Inches(1.3)]
    elif cols == 6:
        widths = [Inches(1.0)] * 6
    else:
        widths = [page_width / cols] * cols

    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.rows[r_idx].cells[c_idx]
            cell.width = widths[c_idx] if c_idx < len(widths) else page_width / cols
            value = row[c_idx] if c_idx < len(row) else ""
            para = cell.paragraphs[0]
            para.text = value
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(2)
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.name = "Calibri"
            if r_idx == 0:
                set_cell_shading(cell)
                for run in para.runs:
                    run.bold = True

    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str]) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)


def add_heading(doc: Document, text: str, level: int) -> None:
    text = strip_md_inline(text)
    if level == 1:
        h = doc.add_heading(text, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_heading(text, level=min(level, 3))


def add_paragraph(doc: Document, text: str) -> None:
    text = strip_md_inline(text)
    if not text:
        return
    if text.startswith("> "):
        text = text[2:]
        para = doc.add_paragraph(text)
        para.paragraph_format.left_indent = Inches(0.35)
        para.runs[0].italic = True
        return
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(6)


def convert() -> None:
    lines = MD_FILE.read_text(encoding="utf-8").splitlines()
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    in_code = False
    code_lines: list[str] = []
    table_rows: list[list[str]] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code blocks
        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Tables
        if stripped.startswith("|") and stripped.endswith("|"):
            if is_table_separator(stripped):
                i += 1
                continue
            table_rows.append(parse_table_row(stripped))
            i += 1
            # Peek ahead: flush table when next line is not a table row
            if i >= len(lines) or not lines[i].strip().startswith("|"):
                add_table(doc, table_rows)
                table_rows = []
            continue

        # Headings
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            add_heading(doc, heading_match.group(2), level)
            i += 1
            continue

        # Horizontal rule / section break
        if stripped in ("---", "***", "___"):
            doc.add_paragraph()
            i += 1
            continue

        # Bullet lists
        if stripped.startswith("- "):
            text = strip_md_inline(stripped[2:])
            doc.add_paragraph(text, style="List Bullet")
            i += 1
            continue

        # Checkbox lists
        if re.match(r"^- \[[ x]\] ", stripped):
            text = strip_md_inline(re.sub(r"^- \[[ x]\] ", "", stripped))
            doc.add_paragraph(f"☐ {text}", style="List Bullet")
            i += 1
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Italic-only footer lines
        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            para = doc.add_paragraph(strip_md_inline(stripped))
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.runs[0].italic = True
            i += 1
            continue

        add_paragraph(doc, stripped)
        i += 1

    doc.save(DOCX_FILE)
    print(f"Created: {DOCX_FILE}")


if __name__ == "__main__":
    convert()
